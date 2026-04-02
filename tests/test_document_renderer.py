"""tests/test_document_renderer.py Document Renderer 测试 for v1.0 Proposal Studio"""
from __future__ import annotations

import os
import tempfile
import pytest

from backend.schemas.proposal_schemas import SectionOutput, ProposalSections
from backend.services.document_renderer import DocumentRenderer


# ─── Fixtures ─────────────────────────────────────────────────────────────────

@pytest.fixture
def sample_proposal_sections() -> ProposalSections:
    """标准的三个 section 的 ProposalSections（无 optional sections）"""
    return ProposalSections(
        workspace_id="a1b2c3d4-e5f6-7890-abcd-ef1234567890",
        pipeline_id="pipeline-001",
        executive_summary=SectionOutput(
            section_key="executive_summary",
            title="一、执行摘要",
            content=(
                "本项目旨在为华道汽车零部件（上海）有限公司提供 JIT 供料解决方案。\n\n"
                "通过部署自动化立体仓库与智能输送系统，预计可将物流效率提升 35%，"
                "年度运营成本节省约 280 万元，投资回收期 2.8 年。"
            ),
            version_id=3,
            tokens_used=512,
            generated_at="2025-01-15T10:00:00Z",
        ),
        core_solution=SectionOutput(
            section_key="core_solution",
            title="二、核心方案设计",
            content=(
                "## 方案概述\n\n本方案采用「自动化立体仓库 + JIT 分拣线」组合模式。\n\n"
                "## 关键设计\n\n"
                "- 仓库面积：35,000 sqm（依据：客户需求 v3）\n"
                "- 日均订单：5,000 单（依据：历史数据 v2）\n"
                "- 人力配置：28 人/班次（依据：工时测算 v3）\n"
                "- 自动化设备：堆垛机 × 4 台、输送线 × 12 套（依据：方案设计 v1）"
            ),
            version_id=3,
            tokens_used=1024,
            generated_at="2025-01-15T10:00:05Z",
        ),
        implementation_plan=SectionOutput(
            section_key="implementation_plan",
            title="三、实施计划",
            content=(
                "## 阶段划分\n\n"
                "1. **需求确认**（第1-2周）：现场调研、数据采集、需求文档评审\n"
                "2. **方案细化**（第3-6周）：详细设计、设备选型、供应商招标\n"
                "3. **实施部署**（第7-20周）：设备安装、系统集成、人员培训\n"
                "4. **验收交付**（第21-24周）：系统联调、试运行、正式验收\n\n"
                "## 关键里程碑\n\n"
                "- M1（第2周末）：需求冻结\n"
                "- M2（第6周末）：方案冻结\n"
                "- M3（第20周末）：设备安装完成\n"
                "- M4（第24周末）：正式验收"
            ),
            version_id=3,
            tokens_used=768,
            generated_at="2025-01-15T10:00:10Z",
        ),
        total_tokens=2304,
        generated_at="2025-01-15T10:00:10Z",
    )


@pytest.fixture
def proposal_with_optional_sections(sample_proposal_sections: ProposalSections) -> ProposalSections:
    """包含 optional sections（financial_kpi + risk_analysis）的 ProposalSections"""
    sample_proposal_sections.financial_kpi = SectionOutput(
        section_key="financial_kpi",
        title="四、财务 KPI",
        content=(
            "## 财务指标摘要\n\n"
            "- 5年 ROI：185%\n"
            "- 投资回收期：2.8 年\n"
            "- 年节省成本：280 万元\n"
            "- NPV（10%折现率）：1,020 万元"
        ),
        version_id=3,
        tokens_used=384,
        generated_at="2025-01-15T10:00:15Z",
    )
    sample_proposal_sections.risk_analysis = SectionOutput(
        section_key="risk_analysis",
        title="五、风险分析",
        content=(
            "## 主要风险及应对\n\n"
            "| 风险 | 概率 | 影响 | 应对措施 |\n"
            "|---|---|---|---|\n"
            "| 设备交付延期 | 中 | 高 | 备选供应商已锁定 |\n"
            "| 人员培训不足 | 低 | 中 | 增加培训周期至 2 周 |\n"
            "| 需求变更 | 高 | 中 | 变更管理流程已建立 |"
        ),
        version_id=3,
        tokens_used=256,
        generated_at="2025-01-15T10:00:20Z",
    )
    return sample_proposal_sections


@pytest.fixture
def sample_active_assumptions() -> list[dict]:
    """标准的 active_assumptions 列表"""
    return [
        {
            "field_key": "warehouse_area",
            "value": "35,000 sqm",
            "rule": "客户需求",
            "version_id": 3,
        },
        {
            "field_key": "daily_orders",
            "value": "5,000 单/天",
            "rule": "历史数据",
            "version_id": 2,
        },
        {
            "field_key": "headcount_per_shift",
            "value": "28 人/班次",
            "rule": "工时测算 v3",
            "version_id": 3,
        },
        {
            "field_key": "automation_level",
            "value": "高",
            "rule": "方案设计 v1",
            "version_id": 1,
        },
    ]


# ─── Tests: render_markdown ───────────────────────────────────────────────────

def test_render_markdown_includes_all_sections(sample_proposal_sections: ProposalSections):
    """验证 Markdown 输出包含三个核心 section（executive_summary / core_solution / implementation_plan）"""
    renderer = DocumentRenderer()
    md = renderer.render_markdown(sample_proposal_sections)

    # 执行摘要
    assert "一、执行摘要" in md
    assert "JIT 供料" in md or "华道汽车" in md

    # 核心方案
    assert "二、核心方案设计" in md
    assert "35,000" in md

    # 实施计划
    assert "三、实施计划" in md
    assert "需求确认" in md

    # 项目编号
    assert "a1b2c3d4" in md


def test_render_markdown_with_optional_sections(proposal_with_optional_sections: ProposalSections):
    """验证当 financial_kpi / risk_analysis 存在时也被渲染"""
    renderer = DocumentRenderer()
    md = renderer.render_markdown(proposal_with_optional_sections)

    assert "四、财务 KPI" in md
    assert "五、风险分析" in md
    assert "ROI" in md
    assert "投资回收期" in md


def test_render_markdown_without_optional_sections(sample_proposal_sections: ProposalSections):
    """验证无 optional sections 时不会报错，且不输出相关内容"""
    renderer = DocumentRenderer()
    md = renderer.render_markdown(sample_proposal_sections)

    assert "财务 KPI" not in md
    assert "风险分析" not in md


def test_render_markdown_includes_assumptions(
    sample_proposal_sections: ProposalSections,
    sample_active_assumptions: list[dict],
):
    """验证附录包含 assumption 清单"""
    renderer = DocumentRenderer()
    md = renderer.render_markdown(sample_proposal_sections, active_assumptions=sample_active_assumptions)

    assert "附录" in md
    assert "warehouse_area" in md
    assert "35,000 sqm" in md
    assert "v3" in md


def test_render_markdown_without_assumptions(sample_proposal_sections: ProposalSections):
    """验证 include_assumptions=False 时不输出附录"""
    renderer = DocumentRenderer()
    md = renderer.render_markdown(sample_proposal_sections, active_assumptions=None)

    assert "附录" not in md


# ─── Tests: render_docx ───────────────────────────────────────────────────────

def test_render_docx_creates_file(sample_proposal_sections: ProposalSections):
    """验证 docx 文件实际生成，文件大小 > 0"""
    renderer = DocumentRenderer()
    with tempfile.TemporaryDirectory() as tmpdir:
        filepath = renderer.render_docx(
            proposal_sections=sample_proposal_sections,
            output_dir=tmpdir,
        )

        assert os.path.exists(filepath), "docx 文件应被创建"
        assert os.path.getsize(filepath) > 0, "docx 文件大小应 > 0"
        assert filepath.endswith(".docx")


def test_render_docx_filename_contains_workspace_id(sample_proposal_sections: ProposalSections):
    """验证生成的文件名包含 workspace_id 前缀"""
    renderer = DocumentRenderer()
    with tempfile.TemporaryDirectory() as tmpdir:
        filepath = renderer.render_docx(
            proposal_sections=sample_proposal_sections,
            output_dir=tmpdir,
        )

        basename = os.path.basename(filepath)
        assert "proposal_a1b2c3d4" in basename


def test_render_docx_includes_assumption_appendix(
    sample_proposal_sections: ProposalSections,
    sample_active_assumptions: list[dict],
):
    """验证附录包含 assumption 清单"""
    renderer = DocumentRenderer()
    with tempfile.TemporaryDirectory() as tmpdir:
        filepath = renderer.render_docx(
            proposal_sections=sample_proposal_sections,
            output_dir=tmpdir,
            include_assumptions=True,
            active_assumptions=sample_active_assumptions,
        )

        # 用 python-docx 重新打开文件，验证附录表格存在
        from docx import Document

        doc = Document(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # 验证附录标题
        assert "附录" in full_text

        # 验证表格存在（至少 1 个表 = 假设清单）
        assert len(doc.tables) >= 1

        # 验证表格内容（需遍历 table cells）
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_texts.append(cell.text)
        table_full = " ".join(table_texts)
        assert "warehouse_area" in table_full, f"warehouse_area not in table: {table_texts}"
        assert "35,000 sqm" in table_full


def test_render_docx_excludes_assumption_appendix_when_disabled(
    sample_proposal_sections: ProposalSections,
    sample_active_assumptions: list[dict],
):
    """验证 include_assumptions=False 时附录不被写入"""
    renderer = DocumentRenderer()
    with tempfile.TemporaryDirectory() as tmpdir:
        filepath = renderer.render_docx(
            proposal_sections=sample_proposal_sections,
            output_dir=tmpdir,
            include_assumptions=False,
            active_assumptions=sample_active_assumptions,
        )

        from docx import Document

        doc = Document(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "附录" not in full_text


def test_render_docx_with_optional_sections(proposal_with_optional_sections: ProposalSections):
    """验证 optional sections（financial_kpi / risk_analysis）被渲染到 docx"""
    renderer = DocumentRenderer()
    with tempfile.TemporaryDirectory() as tmpdir:
        filepath = renderer.render_docx(
            proposal_sections=proposal_with_optional_sections,
            output_dir=tmpdir,
        )

        from docx import Document

        doc = Document(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "财务 KPI" in full_text or "ROI" in full_text
        assert "风险分析" in full_text


def test_render_docx_cover_page(sample_proposal_sections: ProposalSections):
    """验证封面内容（公司名称、日期）"""
    renderer = DocumentRenderer()
    with tempfile.TemporaryDirectory() as tmpdir:
        filepath = renderer.render_docx(
            proposal_sections=sample_proposal_sections,
            output_dir=tmpdir,
            company_name="飞力达物流",
        )

        from docx import Document

        doc = Document(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "飞力达物流" in full_text
        assert "物流解决方案提案书" in full_text


# ─── Tests: edge cases ────────────────────────────────────────────────────────

def test_render_markdown_empty_optional_fields():
    """验证 optional fields 为 None 时不崩溃"""
    renderer = DocumentRenderer()
    sections = ProposalSections(
        workspace_id="test-empty",
        pipeline_id="p1",
        executive_summary=SectionOutput(
            section_key="executive_summary",
            title="一、执行摘要",
            content="test",
            version_id=1,
        ),
        core_solution=SectionOutput(
            section_key="core_solution",
            title="二、核心方案设计",
            content="test",
            version_id=1,
        ),
        implementation_plan=SectionOutput(
            section_key="implementation_plan",
            title="三、实施计划",
            content="test",
            version_id=1,
        ),
    )
    # financial_kpi 和 risk_analysis 保持 None
    md = renderer.render_markdown(sections, active_assumptions=None)
    assert "test" in md
    assert "财务 KPI" not in md
    assert "风险分析" not in md


def test_render_markdown_with_empty_assumptions_list(sample_proposal_sections: ProposalSections):
    """验证空 active_assumptions 列表不崩溃"""
    renderer = DocumentRenderer()
    md = renderer.render_markdown(sample_proposal_sections, active_assumptions=[])
    assert "附录" not in md


def test_render_docx_default_output_dir(sample_proposal_sections: ProposalSections):
    """验证不传 output_dir 时默认写入 /tmp/proposals"""
    renderer = DocumentRenderer()
    filepath = renderer.render_docx(proposal_sections=sample_proposal_sections)
    try:
        assert os.path.exists(filepath)
        assert filepath.startswith("/tmp/proposals/")
    finally:
        # 清理
        if os.path.exists(filepath):
            os.remove(filepath)
