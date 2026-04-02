"""backend/services/document_renderer.py Document Assembly & Export for v1.0 Proposal Studio"""
from __future__ import annotations

import datetime
import os
import uuid
from typing import Optional

from backend.schemas.proposal_schemas import ProposalSections


class DocumentRenderer:
    """
    将 ProposalSections 组装成 Word (.docx) 文档和 Markdown 文本。

    文档结构：
    1. 封面（公司名称 / 项目名称 / 行业 / 日期）
    2. 执行摘要
    3. 核心方案设计
    4. 实施计划
    5. （可选）财务 KPI
    6. （可选）风险分析
    7. 附录：假设清单
    """

    def render_docx(
        self,
        proposal_sections: ProposalSections,
        output_dir: str | None = None,
        company_name: str = "飞力达物流",
        include_assumptions: bool = True,
        active_assumptions: list[dict] | None = None,
    ) -> str:
        """
        将 ProposalSections 渲染为 .docx 文件。

        Args:
            proposal_sections: 三个 section 的内容
            output_dir: 输出目录，默认使用 /tmp/proposals/
            company_name: 封面公司名称
            include_assumptions: 是否在附录中附上 assumption 清单
            active_assumptions: 传给附录的 assumption 列表

        Returns:
            生成的 .docx 文件路径
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── 中文字体设置 helper ──────────────────────────────────────────────

        def set_run_font(run, font_name: str = "微软雅黑", size: int = 11, bold: bool = False):
            """为 Run 设置中文字体（word name + eastAsia）"""
            run.font.name = font_name
            run.font.size = Pt(size)
            run.font.bold = bold
            # w:rPr / w:rFonts / w:eastAsia
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), font_name)
            rPr.insert(0, rFonts)

        def add_heading(doc: Document, text: str, level: int = 1):
            """添加标题段落，使用中文字体"""
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            font_sizes = {1: 18, 2: 14, 3: 12}
            set_run_font(run, font_name="微软雅黑", size=font_sizes.get(level, 11), bold=True)
            return p

        def add_paragraph(doc: Document, text: str):
            """添加正文段落"""
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_run_font(run, font_name="微软雅黑", size=11, bold=False)
            return p

        def add_cover_title(doc: Document, text: str, size: int = 22):
            """居中大标题（封面用）"""
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, font_name="微软雅黑", size=size, bold=True)
            return p

        def add_cover_subtitle(doc: Document, text: str):
            """居中副标题（封面用）"""
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, font_name="微软雅黑", size=13, bold=False)
            return p

        # ── 1. 封面 ──────────────────────────────────────────────────────────
        add_cover_title(doc, company_name, size=22)
        add_cover_title(doc, "物流解决方案提案书", size=16)
        doc.add_paragraph()  # 空行
        add_cover_subtitle(doc, f"项目编号：{proposal_sections.workspace_id[:8]}")
        add_cover_subtitle(
            doc, f"生成时间：{datetime.datetime.now().strftime('%Y年%m月%d日')}"
        )
        doc.add_page_break()

        # ── 2. 执行摘要 ──────────────────────────────────────────────────────
        add_heading(doc, proposal_sections.executive_summary.title, level=1)
        add_paragraph(doc, proposal_sections.executive_summary.content)
        doc.add_page_break()

        # ── 3. 核心方案设计 ──────────────────────────────────────────────────
        add_heading(doc, proposal_sections.core_solution.title, level=1)
        add_paragraph(doc, proposal_sections.core_solution.content)
        doc.add_page_break()

        # ── 4. 实施计划 ──────────────────────────────────────────────────────
        add_heading(doc, proposal_sections.implementation_plan.title, level=1)
        add_paragraph(doc, proposal_sections.implementation_plan.content)

        # ── 5. （可选）财务 KPI ─────────────────────────────────────────────
        if proposal_sections.financial_kpi:
            doc.add_page_break()
            add_heading(doc, proposal_sections.financial_kpi.title, level=1)
            add_paragraph(doc, proposal_sections.financial_kpi.content)

        # ── 6. （可选）风险分析 ─────────────────────────────────────────────
        if proposal_sections.risk_analysis:
            doc.add_page_break()
            add_heading(doc, proposal_sections.risk_analysis.title, level=1)
            add_paragraph(doc, proposal_sections.risk_analysis.content)

        # ── 7. 附录：假设清单 ────────────────────────────────────────────────
        if include_assumptions and active_assumptions:
            doc.add_page_break()
            add_heading(doc, "附录：本方案使用的假设参数", level=1)

            # 表格：字段 | 值 | 依据 | 版本
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            headers = ["参数字段", "假设值", "依据", "版本"]
            for i, hdr in enumerate(headers):
                p = hdr_cells[i].paragraphs[0]
                run = p.add_run(hdr)
                set_run_font(run, font_name="微软雅黑", size=10, bold=True)

            for a in active_assumptions:
                row_cells = table.add_row().cells
                row_cells[0].text = str(a.get("field_key", ""))
                row_cells[1].text = str(a.get("value", ""))
                row_cells[2].text = str(a.get("rule", ""))
                row_cells[3].text = f"v{a.get('version_id', 1)}"

                # 字体设置（表格内）
                for cell in row_cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            set_run_font(run, font_name="微软雅黑", size=10, bold=False)

        # ── 8. 保存 ──────────────────────────────────────────────────────────
        output_dir = output_dir or "/tmp/proposals"
        os.makedirs(output_dir, exist_ok=True)
        filename = (
            f"proposal_{proposal_sections.workspace_id[:8]}"
            f"_{uuid.uuid4().hex[:6]}.docx"
        )
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        return filepath

    def render_markdown(
        self,
        proposal_sections: ProposalSections,
        active_assumptions: list[dict] | None = None,
    ) -> str:
        """
        将 ProposalSections 渲染为 Markdown 文本。

        用于 API 返回或简单预览。

        Args:
            proposal_sections: ProposalSections 对象
            active_assumptions: 可选的 assumption 列表（附录用）

        Returns:
            Markdown 格式的字符串
        """
        lines = []

        # 封面
        lines.append("# 物流解决方案提案书")
        lines.append("")
        lines.append(f"- **项目编号**：`{proposal_sections.workspace_id[:8]}`")
        lines.append(
            f"- **生成时间**：{datetime.datetime.now().strftime('%Y年%m月%d日')}"
        )
        lines.append("")

        # 执行摘要
        lines.append(f"## {proposal_sections.executive_summary.title}")
        lines.append("")
        lines.append(proposal_sections.executive_summary.content)
        lines.append("")

        # 核心方案
        lines.append(f"## {proposal_sections.core_solution.title}")
        lines.append("")
        lines.append(proposal_sections.core_solution.content)
        lines.append("")

        # 实施计划
        lines.append(f"## {proposal_sections.implementation_plan.title}")
        lines.append("")
        lines.append(proposal_sections.implementation_plan.content)
        lines.append("")

        # 财务 KPI（可选）
        if proposal_sections.financial_kpi:
            lines.append(f"## {proposal_sections.financial_kpi.title}")
            lines.append("")
            lines.append(proposal_sections.financial_kpi.content)
            lines.append("")

        # 风险分析（可选）
        if proposal_sections.risk_analysis:
            lines.append(f"## {proposal_sections.risk_analysis.title}")
            lines.append("")
            lines.append(proposal_sections.risk_analysis.content)
            lines.append("")

        # 附录
        if active_assumptions:
            lines.append("## 附录：本方案使用的假设参数")
            lines.append("")
            lines.append("| 参数字段 | 假设值 | 依据 | 版本 |")
            lines.append("|---|---|---|---|")
            for a in active_assumptions:
                field_key = a.get("field_key", "")
                value = a.get("value", "")
                rule = a.get("rule", "")
                version_id = a.get("version_id", 1)
                lines.append(f"| {field_key} | {value} | {rule} | v{version_id} |")
            lines.append("")

        return "\n".join(lines)
